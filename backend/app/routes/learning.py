import json
from datetime import datetime, timezone
from io import BytesIO

from docx import Document
from fastapi import APIRouter, Depends, File, HTTPException, Response, UploadFile
from sqlalchemy import text
from sqlalchemy.orm import Session

from app.database import get_db
from app.models.learning_content import (
    Assessment,
    AssessmentAssignment,
    CaseStudy,
    CaseStudyAssignment,
    CaseStudyDocument,
    CaseStudyPPTTopic,
    CaseStudyVideoTopic,
    ContentBlueprint,
    ContentAssignment,
    IUBreakdown,
    InstructionUnit,
    InstructionalTopic,
    KnowledgeTopic,
    KnowledgeCheck,
    KnowledgeDocument,
    KnowledgePPTTopic,
    KnowledgeVideoTopic,
    LearningActivity,
    LearningProduct,
    MarkingRubric,
    MCQAssessment,
    Module,
    PPTTopic,
    PracticalEvaluation,
    ProductPlan,
    ProjectBrief,
    SkillsActivity,
    VideoTopic,
)
from app.schemas.learning import AdaptiveLearningBlueprintResponse, AssessmentBlueprint, AssessmentProductionBlueprint, ContentBlueprintSchema, InstructionUnitBlueprint, KnowledgeBlueprint, LearningContentChatRequest, LearningContentChatResponse, LearningRequest, LearningResponse, ProductPlanRequest, ProductionTopic, ProjectBriefBlueprint, SkillsBlueprint
from app.services.learning_agent import LIAContentAgent, LearningAgent
from app.services.smtp_service import SMTPService, SMTPServiceError


router = APIRouter(prefix="/learning", tags=["Learning Agent"])
agent = LearningAgent()
content_agent = LIAContentAgent()
smtp_service = SMTPService()


@router.post("/generate", response_model=LearningResponse)
def generate_learning(payload: LearningRequest) -> LearningResponse:
    return agent.generate(payload)


@router.post("/product-plan/generate", response_model=AdaptiveLearningBlueprintResponse)
async def generate_product_plan_blueprint(
    file: UploadFile | None = File(None),
    db: Session = Depends(get_db),
) -> AdaptiveLearningBlueprintResponse:
    _ensure_learning_product_columns(db)
    if not file:
        raise HTTPException(status_code=400, detail="Upload a Product Plan PDF, DOC, or DOCX before generating a blueprint.")

    file_name = ""
    file_name = file.filename or "uploaded-product-plan"
    content = await file.read()
    extracted_text = _extract_product_plan_text(content, file_name)
    metadata = _metadata_from_product_plan(extracted_text, file_name)

    payload = ProductPlanRequest(
        title=metadata["title"],
        module_code=metadata["module_code"],
        course_name=metadata["course_name"],
        learner_level=metadata["learner_level"],
        module_structure=metadata["module_structure"],
        learning_outcomes=metadata["learning_outcomes"],
        total_learning_hours=metadata["total_learning_hours"],
        delivery_modes=metadata["delivery_modes"],
        product_plan_text=extracted_text,
    )
    existing = _find_existing_learning_product(extracted_text, db)
    if existing:
        return _product_response_from_db(existing, db)
    try:
        preview = content_agent.analyze_curriculum(payload)
    except RuntimeError as exc:
        preview = content_agent._build_degraded_curriculum_preview(payload, None, str(exc))
    product = LearningProduct(
        title=preview.title,
        course_name=payload.course_name,
        learner_level=payload.learner_level,
        module_code=payload.module_code,
        product_plan_name=file_name,
        product_plan_text=extracted_text,
        module_structure=payload.module_structure,
        learning_outcomes=payload.learning_outcomes,
        total_learning_hours=preview.total_learning_hours,
        delivery_modes=payload.delivery_modes,
        audience_profile=preview.audience_profile,
        complexity_level="Beginner to Applied",
        readiness_status="Curriculum Review Ready",
    )
    db.add(product)
    db.flush()

    for unit in preview.instruction_units:
        db_unit = InstructionUnit(
            product_id=product.id,
            iu_code=unit.iu_code,
            title=unit.title,
            adaptive_focus=unit.adaptive_focus,
            estimated_hours=unit.estimated_hours,
            complexity_indicator=unit.complexity_indicator,
            learning_goal=unit.learning_goal,
        )
        db.add(db_unit)
        db.flush()
        db.add(
            ContentBlueprint(
                instruction_unit_id=db_unit.id,
                knowledge_instructional_text=json.dumps(unit.blueprint.knowledge_instructional_text),
                ppt_content_topics=json.dumps(unit.blueprint.ppt_content_topics),
                ppt_video_topics=json.dumps(unit.blueprint.ppt_video_topics),
                elearning_activities=json.dumps(unit.blueprint.elearning_activities),
                case_study_text=unit.blueprint.case_study_text,
                case_study_ppt_video_topics=json.dumps(unit.blueprint.case_study_ppt_video_topics),
                case_study_assignments=json.dumps(unit.blueprint.case_study_assignments),
                production_estimate=unit.blueprint.production_estimate,
                readiness_insight=unit.blueprint.readiness_insight,
            )
        )
        db.add(
            Assessment(
                instruction_unit_id=db_unit.id,
                mcqs=json.dumps(unit.assessment.mcqs),
                quizzes=json.dumps(unit.assessment.quizzes),
                focus_areas=json.dumps(unit.assessment.focus_areas),
                assignments=json.dumps(unit.assessment.assignments),
                evaluation_objectives=json.dumps(unit.assessment.evaluation_objectives),
                alignment_check=unit.assessment.alignment_check,
            )
        )

    db.add(
        ProjectBrief(
            product_id=product.id,
            project_brief=preview.project_brief.project_brief,
            capstone_scenario=preview.project_brief.capstone_scenario,
            project_deliverables=json.dumps(preview.project_brief.project_deliverables),
            presentation_outline=json.dumps(preview.project_brief.presentation_outline),
            evaluation_criteria=json.dumps(preview.project_brief.evaluation_criteria),
        )
    )
    db.commit()
    db.refresh(product)
    return preview.model_copy(
        update={
            "id": product.id,
            "module_code": payload.module_code,
            "course_name": payload.course_name,
            "subtitle": payload.course_name or preview.subtitle,
        }
    )


@router.post("/product-plan/{product_id}/iu/{iu_code}/generate", response_model=InstructionUnitBlueprint)
def generate_instruction_unit_blueprint(product_id: int, iu_code: str, db: Session = Depends(get_db)) -> InstructionUnitBlueprint:
    product = db.query(LearningProduct).filter(LearningProduct.id == product_id).first()
    if not product:
        raise HTTPException(status_code=404, detail="Product Plan record not found.")
    cached = _existing_unit_blueprint(product, iu_code, db)
    if cached:
        return cached
    try:
        generated = _generate_unit_blueprint_for_product(product, iu_code, db)
    except RuntimeError as exc:
        raise HTTPException(status_code=502, detail=str(exc)) from exc
    return generated


@router.post("/product-plan/{product_id}/save")
def save_product_plan_review(product_id: int, payload: AdaptiveLearningBlueprintResponse, db: Session = Depends(get_db)) -> dict:
    _ensure_learning_product_columns(db)
    product = db.query(LearningProduct).filter(LearningProduct.id == product_id).first()
    if not product:
        raise HTTPException(status_code=404, detail="Product Plan record not found.")

    product.title = payload.title
    product.course_name = payload.course_name or payload.subtitle or product.course_name
    product.module_code = payload.module_code or product.module_code
    product.total_learning_hours = payload.total_learning_hours
    product.delivery_modes = payload.delivery_modes
    product.audience_profile = payload.audience_profile
    product.readiness_status = "Reviewed and Saved"

    request = ProductPlanRequest(
        title=payload.title,
        module_code=payload.module_code,
        course_name=payload.course_name,
        learner_level=product.learner_level,
        module_structure=product.module_structure,
        learning_outcomes=product.learning_outcomes,
        total_learning_hours=payload.total_learning_hours,
        delivery_modes=payload.delivery_modes,
        product_plan_text=product.product_plan_text,
    )

    for unit in payload.instruction_units:
        stored_unit = (
            db.query(InstructionUnit)
            .filter(InstructionUnit.product_id == product_id, InstructionUnit.iu_code == unit.iu_code)
            .first()
        )
        if not stored_unit:
            stored_unit = InstructionUnit(
                product_id=product_id,
                iu_code=unit.iu_code,
                title=unit.title,
                adaptive_focus=unit.adaptive_focus,
            )
            db.add(stored_unit)
            db.flush()
        stored_unit.title = unit.title
        stored_unit.adaptive_focus = unit.adaptive_focus
        stored_unit.estimated_hours = unit.estimated_hours
        stored_unit.complexity_indicator = unit.complexity_indicator
        stored_unit.learning_goal = unit.learning_goal
        _upsert_unit_outputs(db, stored_unit, unit, request)

    project = db.query(ProjectBrief).filter(ProjectBrief.product_id == product_id).first()
    if not project:
        project = ProjectBrief(product_id=product_id)
        db.add(project)
    project.project_brief = payload.project_brief.project_brief
    project.capstone_scenario = payload.project_brief.capstone_scenario
    project.project_deliverables = json.dumps(payload.project_brief.project_deliverables)
    project.presentation_outline = json.dumps(payload.project_brief.presentation_outline)
    project.evaluation_criteria = json.dumps(payload.project_brief.evaluation_criteria)

    db.commit()
    return {"status": "saved", "product_id": product_id}


@router.get("/content-history")
def get_content_history(db: Session = Depends(get_db)) -> dict:
    _ensure_learning_product_columns(db)
    _ensure_assignment_email_columns(db)
    products = db.query(LearningProduct).order_by(LearningProduct.created_at.asc(), LearningProduct.id.asc()).all()
    records = []
    version_tracker: dict[tuple[str, str], int] = {}
    for product in products:
        module_code = (product.module_code or "").strip() or _extract_module_code(product.product_plan_text or "")
        normalized_title = (product.title or "").strip().lower()
        normalized_code = (module_code or "").strip().upper()
        key = (normalized_title, normalized_code)
        version_tracker[key] = version_tracker.get(key, 0) + 1
        version = version_tracker[key]
        units = db.query(InstructionUnit).filter(InstructionUnit.product_id == product.id).order_by(InstructionUnit.id.asc()).all()
        unit_records = []
        for unit in units:
            assignment_rows = _assignment_rows_for_unit(product, unit, db)
            sent_count = len([row for row in assignment_rows if row.get("email_status") in {"Sent", "Reminder Sent"}])
            unit_records.append(
                {
                    "product_id": product.id,
                    "iu_code": unit.iu_code,
                    "iu_title": unit.title,
                    "learning_goal": unit.learning_goal,
                    "status": "Ready for handoff" if assignment_rows else "Pending generation",
                    "email_status": "Notified" if sent_count else "Not Sent",
                    "topic_count": len(assignment_rows),
                    "assignments": assignment_rows,
                }
            )
        records.append(
            {
                "id": product.id,
                "title": product.title,
                "module_code": module_code,
                "version": version,
                "course_name": product.course_name,
                "readiness_status": product.readiness_status,
                "total_learning_hours": product.total_learning_hours,
                "created_at": product.created_at.isoformat() if product.created_at else "",
                "units": unit_records,
            }
        )
    records.reverse()
    return {"records": records}


@router.post("/content-assignments/save")
def save_content_assignments(payload: dict, db: Session = Depends(get_db)) -> dict:
    _ensure_assignment_email_columns(db)
    rows = payload.get("assignments") if isinstance(payload, dict) else []
    if not isinstance(rows, list):
        raise HTTPException(status_code=400, detail="assignments must be a list.")

    saved = _upsert_assignment_rows(db, rows)
    db.commit()
    return {"status": "saved", "assignments": saved}


@router.post("/content-chat", response_model=LearningContentChatResponse)
def chat_about_learning_content(payload: LearningContentChatRequest, db: Session = Depends(get_db)) -> LearningContentChatResponse:
    question = (payload.question or "").strip()
    if not question:
        raise HTTPException(status_code=400, detail="Please enter a question for LIA Content Chat.")
    product = db.query(LearningProduct).filter(LearningProduct.id == payload.product_id).first()
    if not product:
        raise HTTPException(status_code=404, detail="Product Plan record not found.")
    try:
        return LearningContentChatResponse(**_content_chat_response(product, question, db))
    except RuntimeError as exc:
        raise HTTPException(status_code=502, detail=str(exc)) from exc


@router.post("/content-assignments/notify")
def notify_content_assignments(payload: dict, db: Session = Depends(get_db)) -> dict:
    _ensure_assignment_email_columns(db)
    product, unit = _assignment_product_unit(payload, db)
    rows = payload.get("assignments") if isinstance(payload, dict) else []
    if isinstance(rows, list) and rows:
        _upsert_assignment_rows(db, rows)
        db.commit()

    assignments = _assignment_query(product.id, unit.iu_code, db)
    missing_deadlines = [row.topic_code for row in assignments if not (row.deadline or "").strip()]
    if missing_deadlines:
        raise HTTPException(status_code=400, detail="Set a deadline for all IU tasks before sending handoff.")
    missing_owner = [row.topic_code for row in assignments if not (row.owner or "").strip()]
    if missing_owner:
        raise HTTPException(status_code=400, detail="Assign an owner for all IU tasks before sending handoff.")
    recipients = _assignment_recipients(payload.get("recipients") if isinstance(payload, dict) else None, assignments)
    if not recipients:
        raise HTTPException(status_code=400, detail="Add at least one assignee email before sending the IU handoff.")

    selected_format = str(payload.get("format") or "docx").lower().strip(".")
    attachment = _unit_download_attachment(product, unit, selected_format, db)
    subject = f"LIA Content Production Handoff: {unit.iu_code} - {unit.title}"
    body = _assignment_email_body(product, unit, assignments, reminder=False)
    try:
        smtp_service.send(recipients, subject, body, attachments=[attachment])
    except SMTPServiceError as exc:
        raise HTTPException(status_code=exc.status_code, detail=exc.message) from exc

    timestamp = _now_stamp()
    for assignment in assignments:
        assignment.email_status = "Sent"
        assignment.email_sent_at = timestamp
    db.commit()
    return {"status": "sent", "message": "IU handoff email sent successfully.", "recipients": recipients}


@router.delete("/content-history")
def clear_content_history(db: Session = Depends(get_db)) -> dict:
    _ensure_assignment_email_columns(db)
    db.query(ContentAssignment).delete()
    db.query(MarkingRubric).delete()
    db.query(AssessmentAssignment).delete()
    db.query(MCQAssessment).delete()
    db.query(CaseStudyAssignment).delete()
    db.query(CaseStudyVideoTopic).delete()
    db.query(CaseStudyPPTTopic).delete()
    db.query(CaseStudyDocument).delete()
    db.query(LearningActivity).delete()
    db.query(KnowledgeVideoTopic).delete()
    db.query(KnowledgePPTTopic).delete()
    db.query(KnowledgeDocument).delete()
    db.query(InstructionalTopic).delete()
    db.query(IUBreakdown).delete()
    db.query(ProjectBrief).delete()
    db.query(Assessment).delete()
    db.query(ContentBlueprint).delete()
    db.query(InstructionUnit).delete()
    db.query(LearningProduct).delete()
    db.commit()
    return {"status": "cleared"}


@router.delete("/content-history/{product_id}")
def clear_content_history_record(product_id: int, db: Session = Depends(get_db)) -> dict:
    _ensure_assignment_email_columns(db)
    product = db.query(LearningProduct).filter(LearningProduct.id == product_id).first()
    if not product:
        raise HTTPException(status_code=404, detail="History record not found.")
    unit_ids = [row[0] for row in db.query(InstructionUnit.id).filter(InstructionUnit.product_id == product_id).all()]
    if unit_ids:
        db.query(ContentAssignment).filter(ContentAssignment.product_id == product_id).delete(synchronize_session=False)
        db.query(MarkingRubric).filter(MarkingRubric.instruction_unit_id.in_(unit_ids)).delete(synchronize_session=False)
        db.query(AssessmentAssignment).filter(AssessmentAssignment.instruction_unit_id.in_(unit_ids)).delete(synchronize_session=False)
        db.query(MCQAssessment).filter(MCQAssessment.instruction_unit_id.in_(unit_ids)).delete(synchronize_session=False)
        db.query(CaseStudyAssignment).filter(CaseStudyAssignment.instruction_unit_id.in_(unit_ids)).delete(synchronize_session=False)
        db.query(CaseStudyVideoTopic).filter(CaseStudyVideoTopic.instruction_unit_id.in_(unit_ids)).delete(synchronize_session=False)
        db.query(CaseStudyPPTTopic).filter(CaseStudyPPTTopic.instruction_unit_id.in_(unit_ids)).delete(synchronize_session=False)
        db.query(CaseStudyDocument).filter(CaseStudyDocument.instruction_unit_id.in_(unit_ids)).delete(synchronize_session=False)
        db.query(LearningActivity).filter(LearningActivity.instruction_unit_id.in_(unit_ids)).delete(synchronize_session=False)
        db.query(KnowledgeVideoTopic).filter(KnowledgeVideoTopic.instruction_unit_id.in_(unit_ids)).delete(synchronize_session=False)
        db.query(KnowledgePPTTopic).filter(KnowledgePPTTopic.instruction_unit_id.in_(unit_ids)).delete(synchronize_session=False)
        db.query(KnowledgeDocument).filter(KnowledgeDocument.instruction_unit_id.in_(unit_ids)).delete(synchronize_session=False)
        db.query(InstructionalTopic).filter(InstructionalTopic.instruction_unit_id.in_(unit_ids)).delete(synchronize_session=False)
        db.query(IUBreakdown).filter(IUBreakdown.instruction_unit_id.in_(unit_ids)).delete(synchronize_session=False)
        db.query(Assessment).filter(Assessment.instruction_unit_id.in_(unit_ids)).delete(synchronize_session=False)
        db.query(ContentBlueprint).filter(ContentBlueprint.instruction_unit_id.in_(unit_ids)).delete(synchronize_session=False)
        db.query(InstructionUnit).filter(InstructionUnit.id.in_(unit_ids)).delete(synchronize_session=False)
    db.query(ProjectBrief).filter(ProjectBrief.product_id == product_id).delete(synchronize_session=False)
    db.query(LearningProduct).filter(LearningProduct.id == product_id).delete(synchronize_session=False)
    db.commit()
    return {"status": "cleared", "product_id": product_id}


@router.post("/content-history/{product_id}/clear")
def clear_content_history_record_post(product_id: int, db: Session = Depends(get_db)) -> dict:
    return clear_content_history_record(product_id, db)


@router.post("/content-assignments/remind")
def remind_content_assignments(payload: dict, db: Session = Depends(get_db)) -> dict:
    _ensure_assignment_email_columns(db)
    product, unit = _assignment_product_unit(payload, db)
    rows = payload.get("assignments") if isinstance(payload, dict) else []
    if isinstance(rows, list) and rows:
        _upsert_assignment_rows(db, rows)
        db.commit()

    assignments = [row for row in _assignment_query(product.id, unit.iu_code, db) if row.status != "Completed"]
    recipients = _assignment_recipients(payload.get("recipients") if isinstance(payload, dict) else None, assignments)
    if not recipients:
        raise HTTPException(status_code=400, detail="Add assignee emails for incomplete tasks before sending reminders.")

    subject = f"Reminder: {unit.iu_code} content production tasks need follow-up"
    body = _assignment_email_body(product, unit, assignments, reminder=True)
    try:
        smtp_service.send(recipients, subject, body)
    except SMTPServiceError as exc:
        raise HTTPException(status_code=exc.status_code, detail=exc.message) from exc

    timestamp = _now_stamp()
    for assignment in assignments:
        assignment.email_status = "Reminder Sent"
        assignment.reminder_sent_at = timestamp
    db.commit()
    return {"status": "sent", "message": "Reminder email sent successfully.", "recipients": recipients}


@router.get("/product-plan/{product_id}/project-brief/download")
def download_project_brief(product_id: int, format: str = "docx", db: Session = Depends(get_db)) -> Response:
    product = db.query(LearningProduct).filter(LearningProduct.id == product_id).first()
    if not product:
        raise HTTPException(status_code=404, detail="Product Plan record not found.")

    content = _project_brief_download_text(product, db)
    safe_name = f"{_download_safe_name(product.title)}-Project-Brief"
    selected = format.lower().strip(".")
    if selected == "docx":
        binary = _project_brief_docx(content)
        return Response(
            content=binary,
            media_type="application/vnd.openxmlformats-officedocument.wordprocessingml.document",
            headers={"Content-Disposition": f'attachment; filename="{safe_name}.docx"'},
        )
    if selected in {"ppt", "pptx"}:
        html = _project_brief_ppt(content)
        return Response(
            content=html,
            media_type="application/vnd.ms-powerpoint",
            headers={"Content-Disposition": f'attachment; filename="{safe_name}.ppt"'},
        )
    raise HTTPException(status_code=400, detail="Supported formats are docx and ppt.")


@router.get("/product-plan/{product_id}/iu/{iu_code}/download")
def download_instruction_unit_content(product_id: int, iu_code: str, format: str = "docx", db: Session = Depends(get_db)) -> Response:
    product = db.query(LearningProduct).filter(LearningProduct.id == product_id).first()
    if not product:
        raise HTTPException(status_code=404, detail="Product Plan record not found.")
    unit = (
        db.query(InstructionUnit)
        .filter(InstructionUnit.product_id == product_id, InstructionUnit.iu_code == iu_code.upper())
        .first()
    )
    if not unit:
        raise HTTPException(status_code=404, detail=f"{iu_code.upper()} was not found for this Product Plan.")

    content = _unit_download_text(product, unit, db)
    safe_name = f"{unit.iu_code}-{_download_safe_name(unit.title)}"
    selected = format.lower().strip(".")
    if selected == "docx":
        binary = _unit_docx(content)
        return Response(
            content=binary,
            media_type="application/vnd.openxmlformats-officedocument.wordprocessingml.document",
            headers={"Content-Disposition": f'attachment; filename="{safe_name}.docx"'},
        )
    if selected in {"xls", "xlsx"}:
        html = _unit_xls(content, unit)
        return Response(
            content=html,
            media_type="application/vnd.ms-excel",
            headers={"Content-Disposition": f'attachment; filename="{safe_name}.xls"'},
        )
    if selected == "pdf":
        binary = _simple_pdf(content)
        return Response(
            content=binary,
            media_type="application/pdf",
            headers={"Content-Disposition": f'attachment; filename="{safe_name}.pdf"'},
        )
    return Response(
        content=content,
        media_type="text/plain",
        headers={"Content-Disposition": f'attachment; filename="{safe_name}.txt"'},
    )


@router.post("/product-plan/generate-json", response_model=AdaptiveLearningBlueprintResponse)
def generate_product_plan_blueprint_json(payload: ProductPlanRequest, db: Session = Depends(get_db)) -> AdaptiveLearningBlueprintResponse:
    existing = _find_existing_learning_product(payload.product_plan_text, db)
    if existing:
        return _product_response_from_db(existing, db)
    return _persist_blueprint(payload, "", db)


def _persist_blueprint(payload: ProductPlanRequest, file_name: str, db: Session) -> AdaptiveLearningBlueprintResponse:
    _ensure_learning_product_columns(db)
    preview = content_agent.generate_blueprint(payload)
    product = LearningProduct(
        title=preview.title,
        course_name=payload.course_name,
        learner_level=payload.learner_level,
        module_code=payload.module_code,
        product_plan_name=file_name,
        product_plan_text=payload.product_plan_text,
        module_structure=payload.module_structure,
        learning_outcomes=payload.learning_outcomes,
        total_learning_hours=preview.total_learning_hours,
        delivery_modes=payload.delivery_modes,
        audience_profile=preview.audience_profile,
        complexity_level="Beginner to Applied",
        readiness_status="Production Blueprint Ready",
    )
    db.add(product)
    db.flush()

    for unit in preview.instruction_units:
        db_unit = InstructionUnit(
            product_id=product.id,
            iu_code=unit.iu_code,
            title=unit.title,
            adaptive_focus=unit.adaptive_focus,
            estimated_hours=unit.estimated_hours,
            complexity_indicator=unit.complexity_indicator,
            learning_goal=unit.learning_goal,
        )
        db.add(db_unit)
        db.flush()
        db.add(
            ContentBlueprint(
                instruction_unit_id=db_unit.id,
                knowledge_instructional_text=json.dumps(unit.blueprint.knowledge_instructional_text),
                ppt_content_topics=json.dumps(unit.blueprint.ppt_content_topics),
                ppt_video_topics=json.dumps(unit.blueprint.ppt_video_topics),
                elearning_activities=json.dumps(unit.blueprint.elearning_activities),
                case_study_text=unit.blueprint.case_study_text,
                case_study_ppt_video_topics=json.dumps(unit.blueprint.case_study_ppt_video_topics),
                case_study_assignments=json.dumps(unit.blueprint.case_study_assignments),
                production_estimate=unit.blueprint.production_estimate,
                readiness_insight=unit.blueprint.readiness_insight,
            )
        )
        db.add(
            Assessment(
                instruction_unit_id=db_unit.id,
                mcqs=json.dumps(unit.assessment.mcqs),
                quizzes=json.dumps(unit.assessment.quizzes),
                focus_areas=json.dumps(unit.assessment.focus_areas),
                assignments=json.dumps(unit.assessment.assignments),
                evaluation_objectives=json.dumps(unit.assessment.evaluation_objectives),
                alignment_check=unit.assessment.alignment_check,
            )
        )

    db.add(
        ProjectBrief(
            product_id=product.id,
            project_brief=preview.project_brief.project_brief,
            capstone_scenario=preview.project_brief.capstone_scenario,
            project_deliverables=json.dumps(preview.project_brief.project_deliverables),
            presentation_outline=json.dumps(preview.project_brief.presentation_outline),
            evaluation_criteria=json.dumps(preview.project_brief.evaluation_criteria),
        )
    )
    db.commit()
    db.refresh(product)
    _persist_repository_tables(payload, preview, product.id, file_name, db)
    return preview.model_copy(
        update={
            "id": product.id,
            "module_code": payload.module_code,
            "course_name": payload.course_name,
            "subtitle": payload.course_name or preview.subtitle,
        }
    )


def _find_existing_learning_product(product_plan_text: str, db: Session) -> LearningProduct | None:
    normalized = _normalize_lookup_text(product_plan_text)
    if not normalized:
        return None
    products = db.query(LearningProduct).order_by(LearningProduct.created_at.desc(), LearningProduct.id.desc()).all()
    for product in products:
        if _normalize_lookup_text(product.product_plan_text or "") == normalized:
            return product
    return None


def _product_response_from_db(product: LearningProduct, db: Session) -> AdaptiveLearningBlueprintResponse:
    units = db.query(InstructionUnit).filter(InstructionUnit.product_id == product.id).order_by(InstructionUnit.id.asc()).all()
    project = db.query(ProjectBrief).filter(ProjectBrief.product_id == product.id).first()
    return AdaptiveLearningBlueprintResponse(
        id=product.id,
        title=product.title,
        subtitle=product.course_name or "Adaptive Learning Content Intelligence Platform",
        module_code=product.module_code or "",
        course_name=product.course_name or "",
        audience_profile=product.audience_profile or "",
        total_learning_hours=product.total_learning_hours,
        delivery_modes=product.delivery_modes or "",
        curriculum_analysis=[f"Retrieved existing Product Plan blueprint from database for {product.title}."],
        adaptive_learning_recommendations=["Using saved curriculum intelligence to avoid repeat AI generation during the demo."],
        complexity_indicators=[product.complexity_level or "Beginner to Applied"],
        assessment_alignment_checks=["Assessment alignment retrieved from saved IU records."],
        content_production_estimation=["Production estimates loaded from saved blueprint data."],
        delivery_readiness_insights=[product.readiness_status or "Reviewed and Saved"],
        instruction_units=[_unit_blueprint_from_db(unit, product, db) for unit in units],
        project_brief=ProjectBriefBlueprint(
            project_brief=project.project_brief if project else "",
            capstone_scenario=project.capstone_scenario if project else "",
            project_deliverables=_json_load_list(project.project_deliverables if project else "[]"),
            presentation_outline=_json_load_list(project.presentation_outline if project else "[]"),
            evaluation_criteria=_json_load_list(project.evaluation_criteria if project else "[]"),
        ),
    )


def _existing_unit_blueprint(product: LearningProduct, iu_code: str, db: Session) -> InstructionUnitBlueprint | None:
    unit = (
        db.query(InstructionUnit)
        .filter(InstructionUnit.product_id == product.id, InstructionUnit.iu_code == iu_code.upper())
        .first()
    )
    if not unit:
        return None
    if not _unit_has_saved_content(unit, db):
        return None
    return _unit_blueprint_from_db(unit, product, db)


def _unit_has_saved_content(unit: InstructionUnit, db: Session) -> bool:
    return any(
        (
            db.query(KnowledgeDocument).filter(KnowledgeDocument.instruction_unit_id == unit.id).first(),
            db.query(CaseStudyAssignment).filter(CaseStudyAssignment.instruction_unit_id == unit.id).first(),
            db.query(MarkingRubric).filter(MarkingRubric.instruction_unit_id == unit.id).first(),
        )
    )


def _unit_blueprint_from_db(unit: InstructionUnit, product: LearningProduct, db: Session) -> InstructionUnitBlueprint:
    breakdown = db.query(IUBreakdown).filter(IUBreakdown.instruction_unit_id == unit.id).order_by(IUBreakdown.id.desc()).first()
    blueprint = db.query(ContentBlueprint).filter(ContentBlueprint.instruction_unit_id == unit.id).first()
    assessment = db.query(Assessment).filter(Assessment.instruction_unit_id == unit.id).first()
    knowledge_docs = db.query(KnowledgeDocument).filter(KnowledgeDocument.instruction_unit_id == unit.id).all()
    knowledge_ppt = db.query(KnowledgePPTTopic).filter(KnowledgePPTTopic.instruction_unit_id == unit.id).all()
    knowledge_video = db.query(KnowledgeVideoTopic).filter(KnowledgeVideoTopic.instruction_unit_id == unit.id).all()
    learning_activities = db.query(LearningActivity).filter(LearningActivity.instruction_unit_id == unit.id).all()
    case_docs = db.query(CaseStudyDocument).filter(CaseStudyDocument.instruction_unit_id == unit.id).all()
    case_ppt = db.query(CaseStudyPPTTopic).filter(CaseStudyPPTTopic.instruction_unit_id == unit.id).all()
    case_video = db.query(CaseStudyVideoTopic).filter(CaseStudyVideoTopic.instruction_unit_id == unit.id).all()
    case_assignments = db.query(CaseStudyAssignment).filter(CaseStudyAssignment.instruction_unit_id == unit.id).all()
    mcqs = db.query(MCQAssessment).filter(MCQAssessment.instruction_unit_id == unit.id).all()
    assessment_assignments = db.query(AssessmentAssignment).filter(AssessmentAssignment.instruction_unit_id == unit.id).all()
    rubrics = db.query(MarkingRubric).filter(MarkingRubric.instruction_unit_id == unit.id).all()

    return InstructionUnitBlueprint(
        iu_code=unit.iu_code,
        title=unit.title,
        module_code=(breakdown.module_code if breakdown else product.module_code) or "",
        module_name=(breakdown.module_name if breakdown else product.title) or product.title,
        adaptive_focus=unit.adaptive_focus,
        estimated_hours=unit.estimated_hours,
        complexity_indicator=unit.complexity_indicator,
        delivery_mode=breakdown.delivery_mode if breakdown else "Self-paced document-led adaptive learning",
        learning_goal=unit.learning_goal,
        knowledge=KnowledgeBlueprint(
            instructional_content_text=_topic_models(knowledge_docs, "topic_code", "topic_title"),
            ppt_text=_topic_models(knowledge_ppt, "topic_code", "topic_title"),
            ppt_videos_podcast=_topic_models(knowledge_video, "topic_code", "topic_title"),
            e_learning=_topic_models(learning_activities, "activity_code", "activity_title"),
        ),
        skills=SkillsBlueprint(
            case_study_word_document=_topic_models(case_docs, "case_code", "case_title"),
            case_study_ppt=_topic_models(case_ppt, "topic_code", "topic_title"),
            case_study_demo_videos=_topic_models(case_video, "topic_code", "topic_title"),
            case_study_assignment=_topic_models(case_assignments, "assignment_code", "assignment_title"),
        ),
        assessment_blueprint=AssessmentProductionBlueprint(
            mcq_assessment=_topic_models(mcqs, "assessment_code", "assessment_title"),
            assessment_assignment=_topic_models(assessment_assignments, "assignment_code", "assignment_title"),
            marking_rubrics=_topic_models(rubrics, "rubric_code", "rubric_title"),
        ),
        blueprint=ContentBlueprintSchema(
            knowledge_instructional_text=_json_load_list(blueprint.knowledge_instructional_text if blueprint else "[]"),
            ppt_content_topics=_json_load_list(blueprint.ppt_content_topics if blueprint else "[]"),
            ppt_video_topics=_json_load_list(blueprint.ppt_video_topics if blueprint else "[]"),
            elearning_activities=_json_load_list(blueprint.elearning_activities if blueprint else "[]"),
            case_study_text=blueprint.case_study_text if blueprint else "",
            case_study_ppt_video_topics=_json_load_list(blueprint.case_study_ppt_video_topics if blueprint else "[]"),
            case_study_assignments=_json_load_list(blueprint.case_study_assignments if blueprint else "[]"),
            learning_sequence=[],
            glossary_concepts=[],
            adaptive_learning_support=[],
            learning_outcome_alignment=[],
            production_estimate=blueprint.production_estimate if blueprint else "",
            readiness_insight=blueprint.readiness_insight if blueprint else "",
        ),
        assessment=AssessmentBlueprint(
            mcqs=_json_load_list(assessment.mcqs if assessment else "[]"),
            focus_areas=_json_load_list(assessment.focus_areas if assessment else "[]"),
            quizzes=_json_load_list(assessment.quizzes if assessment else "[]"),
            assignments=_json_load_list(assessment.assignments if assessment else "[]"),
            evaluation_objectives=_json_load_list(assessment.evaluation_objectives if assessment else "[]"),
            alignment_check=assessment.alignment_check if assessment else "",
        ),
    )


def _topic_models(rows: list, code_attr: str, title_attr: str) -> list[ProductionTopic]:
    return [
        ProductionTopic(
            code=str(getattr(row, code_attr, "") or ""),
            title=str(getattr(row, title_attr, "") or ""),
            description=str(getattr(row, "description", "") or ""),
        )
        for row in rows
    ]


def _upsert_unit_outputs(db: Session, stored_unit: InstructionUnit, unit: InstructionUnitBlueprint, payload: ProductPlanRequest) -> None:
    db.query(IUBreakdown).filter(IUBreakdown.instruction_unit_id == stored_unit.id).delete()
    db.query(InstructionalTopic).filter(InstructionalTopic.instruction_unit_id == stored_unit.id).delete()
    db.query(KnowledgeDocument).filter(KnowledgeDocument.instruction_unit_id == stored_unit.id).delete()
    db.query(KnowledgePPTTopic).filter(KnowledgePPTTopic.instruction_unit_id == stored_unit.id).delete()
    db.query(KnowledgeVideoTopic).filter(KnowledgeVideoTopic.instruction_unit_id == stored_unit.id).delete()
    db.query(LearningActivity).filter(LearningActivity.instruction_unit_id == stored_unit.id).delete()
    db.query(CaseStudyDocument).filter(CaseStudyDocument.instruction_unit_id == stored_unit.id).delete()
    db.query(CaseStudyPPTTopic).filter(CaseStudyPPTTopic.instruction_unit_id == stored_unit.id).delete()
    db.query(CaseStudyVideoTopic).filter(CaseStudyVideoTopic.instruction_unit_id == stored_unit.id).delete()
    db.query(CaseStudyAssignment).filter(CaseStudyAssignment.instruction_unit_id == stored_unit.id).delete()
    db.query(MCQAssessment).filter(MCQAssessment.instruction_unit_id == stored_unit.id).delete()
    db.query(AssessmentAssignment).filter(AssessmentAssignment.instruction_unit_id == stored_unit.id).delete()
    db.query(MarkingRubric).filter(MarkingRubric.instruction_unit_id == stored_unit.id).delete()

    blueprint = db.query(ContentBlueprint).filter(ContentBlueprint.instruction_unit_id == stored_unit.id).first()
    if not blueprint:
        blueprint = ContentBlueprint(instruction_unit_id=stored_unit.id)
        db.add(blueprint)
    blueprint.knowledge_instructional_text = json.dumps(unit.blueprint.knowledge_instructional_text)
    blueprint.ppt_content_topics = json.dumps(unit.blueprint.ppt_content_topics)
    blueprint.ppt_video_topics = json.dumps(unit.blueprint.ppt_video_topics)
    blueprint.elearning_activities = json.dumps(unit.blueprint.elearning_activities)
    blueprint.case_study_text = unit.blueprint.case_study_text
    blueprint.case_study_ppt_video_topics = json.dumps(unit.blueprint.case_study_ppt_video_topics)
    blueprint.case_study_assignments = json.dumps(unit.blueprint.case_study_assignments)
    blueprint.production_estimate = unit.blueprint.production_estimate
    blueprint.readiness_insight = unit.blueprint.readiness_insight

    assessment = db.query(Assessment).filter(Assessment.instruction_unit_id == stored_unit.id).first()
    if not assessment:
        assessment = Assessment(instruction_unit_id=stored_unit.id)
        db.add(assessment)
    assessment.mcqs = json.dumps(unit.assessment.mcqs)
    assessment.quizzes = json.dumps(unit.assessment.quizzes)
    assessment.focus_areas = json.dumps(unit.assessment.focus_areas)
    assessment.assignments = json.dumps(unit.assessment.assignments)
    assessment.evaluation_objectives = json.dumps(unit.assessment.evaluation_objectives)
    assessment.alignment_check = unit.assessment.alignment_check

    db.add(
        IUBreakdown(
            instruction_unit_id=stored_unit.id,
            iu_code=unit.iu_code,
            iu_title=unit.title,
            module_code=unit.module_code,
            module_name=unit.module_name,
            learning_goal=unit.learning_goal,
            learner_level=payload.learner_level,
            delivery_mode=unit.delivery_mode,
            complexity_indicator=unit.complexity_indicator,
            estimated_hours=unit.estimated_hours,
        )
    )
    for topic in unit.knowledge.instructional_content_text:
        db.add(KnowledgeDocument(instruction_unit_id=stored_unit.id, topic_code=topic.code, topic_title=topic.title, description=topic.description))
        db.add(InstructionalTopic(instruction_unit_id=stored_unit.id, topic_code=topic.code, topic_title=topic.title, description=topic.description, section="Knowledge", asset_type="Instructional Content Text"))
    for topic in unit.knowledge.e_learning:
        db.add(LearningActivity(instruction_unit_id=stored_unit.id, activity_code=topic.code, activity_title=topic.title, description=topic.description, activity_type="E-learning"))
    for topic in unit.knowledge.ppt_text:
        db.add(KnowledgePPTTopic(instruction_unit_id=stored_unit.id, topic_code=topic.code, topic_title=topic.title, description=topic.description))
    for topic in unit.knowledge.ppt_videos_podcast:
        db.add(KnowledgeVideoTopic(instruction_unit_id=stored_unit.id, topic_code=topic.code, topic_title=topic.title, description=topic.description))
    for topic in unit.skills.case_study_word_document:
        db.add(CaseStudyDocument(instruction_unit_id=stored_unit.id, case_code=topic.code, case_title=topic.title, description=topic.description))
    for topic in unit.skills.case_study_ppt:
        db.add(CaseStudyPPTTopic(instruction_unit_id=stored_unit.id, topic_code=topic.code, topic_title=topic.title, description=topic.description))
    for topic in unit.skills.case_study_demo_videos:
        db.add(CaseStudyVideoTopic(instruction_unit_id=stored_unit.id, topic_code=topic.code, topic_title=topic.title, description=topic.description))
    for topic in unit.skills.case_study_assignment:
        db.add(CaseStudyAssignment(instruction_unit_id=stored_unit.id, assignment_code=topic.code, assignment_title=topic.title, description=topic.description))
    for topic in unit.assessment_blueprint.mcq_assessment:
        db.add(MCQAssessment(instruction_unit_id=stored_unit.id, assessment_code=topic.code, assessment_title=topic.title, description=topic.description))
    for topic in unit.assessment_blueprint.assessment_assignment:
        db.add(AssessmentAssignment(instruction_unit_id=stored_unit.id, assignment_code=topic.code, assignment_title=topic.title, description=topic.description))
    for topic in unit.assessment_blueprint.marking_rubrics:
        db.add(MarkingRubric(instruction_unit_id=stored_unit.id, rubric_code=topic.code, rubric_title=topic.title, description=topic.description))


def _generate_unit_blueprint_for_product(product: LearningProduct, iu_code: str, db: Session) -> InstructionUnitBlueprint:
    normalized_code = iu_code.upper()
    unit_row = (
        db.query(InstructionUnit)
        .filter(InstructionUnit.product_id == product.id, InstructionUnit.iu_code == normalized_code)
        .first()
    )
    if not unit_row:
        raise HTTPException(status_code=404, detail=f"{normalized_code} was not found for this Product Plan.")

    payload = ProductPlanRequest(
        title=product.title,
        module_code=_extract_module_code(product.product_plan_text),
        course_name=product.course_name,
        learner_level=product.learner_level,
        module_structure=product.module_structure,
        learning_outcomes=product.learning_outcomes,
        total_learning_hours=product.total_learning_hours,
        delivery_modes=product.delivery_modes,
        product_plan_text=product.product_plan_text,
    )
    unit = content_agent._empty_iu_blueprint(
        code=unit_row.iu_code,
        title=unit_row.title,
        focus=unit_row.adaptive_focus,
        complexity=unit_row.complexity_indicator,
        hours=unit_row.estimated_hours,
        outcome=unit_row.learning_goal,
        payload=payload,
    )
    generated = content_agent.generate_iu_blueprint(payload, unit)
    unit_row.title = generated.title
    unit_row.adaptive_focus = generated.adaptive_focus
    unit_row.estimated_hours = generated.estimated_hours
    unit_row.complexity_indicator = generated.complexity_indicator
    unit_row.learning_goal = generated.learning_goal
    _upsert_unit_outputs(db, unit_row, generated, payload)
    db.commit()
    return generated


def _persist_repository_tables(payload: ProductPlanRequest, blueprint: AdaptiveLearningBlueprintResponse, product_id: int, file_name: str, db: Session) -> None:
    product_plan = ProductPlan(
        course_name=payload.course_name,
        learner_level=payload.learner_level,
        source_filename=file_name,
        extracted_text=payload.product_plan_text,
        learning_outcomes=payload.learning_outcomes,
        delivery_modes=payload.delivery_modes,
        total_learning_hours=blueprint.total_learning_hours,
    )
    db.add(product_plan)
    db.flush()
    module = Module(
        product_plan_id=product_plan.id,
        module_code=blueprint.module_code or payload.module_code,
        module_name=blueprint.title,
        module_hours=blueprint.total_learning_hours,
    )
    db.add(module)
    db.flush()
    stored_units = db.query(InstructionUnit).filter(InstructionUnit.product_id == product_id).order_by(InstructionUnit.id.asc()).all()
    for stored_unit, unit in zip(stored_units, blueprint.instruction_units):
        db.add(
            IUBreakdown(
                instruction_unit_id=stored_unit.id,
                iu_code=unit.iu_code,
                iu_title=unit.title,
                module_code=unit.module_code,
                module_name=unit.module_name,
                learning_goal=unit.learning_goal,
                learner_level=payload.learner_level,
                delivery_mode=unit.delivery_mode,
                complexity_indicator=unit.complexity_indicator,
                estimated_hours=unit.estimated_hours,
            )
        )
        for topic in unit.knowledge.instructional_content_text:
            db.add(KnowledgeTopic(instruction_unit_id=stored_unit.id, topic_code=topic.code, topic_title=topic.title, description=topic.description, asset_type="Instructional Content Text"))
            db.add(KnowledgeDocument(instruction_unit_id=stored_unit.id, topic_code=topic.code, topic_title=topic.title, description=topic.description))
            db.add(InstructionalTopic(instruction_unit_id=stored_unit.id, topic_code=topic.code, topic_title=topic.title, description=topic.description, section="Knowledge", asset_type="Instructional Content Text"))
        for topic in unit.knowledge.e_learning:
            db.add(KnowledgeTopic(instruction_unit_id=stored_unit.id, topic_code=topic.code, topic_title=topic.title, description=topic.description, asset_type="E-learning"))
            db.add(LearningActivity(instruction_unit_id=stored_unit.id, activity_code=topic.code, activity_title=topic.title, description=topic.description, activity_type="E-learning"))
        for topic in unit.knowledge.ppt_text:
            db.add(PPTTopic(instruction_unit_id=stored_unit.id, topic_code=topic.code, topic_title=topic.title, description=topic.description))
            db.add(KnowledgePPTTopic(instruction_unit_id=stored_unit.id, topic_code=topic.code, topic_title=topic.title, description=topic.description))
            db.add(InstructionalTopic(instruction_unit_id=stored_unit.id, topic_code=topic.code, topic_title=topic.title, description=topic.description, section="Knowledge", asset_type="PPT Text"))
        for topic in unit.knowledge.ppt_videos_podcast:
            db.add(VideoTopic(instruction_unit_id=stored_unit.id, topic_code=topic.code, topic_title=topic.title, description=topic.description))
            db.add(KnowledgeVideoTopic(instruction_unit_id=stored_unit.id, topic_code=topic.code, topic_title=topic.title, description=topic.description))
            db.add(InstructionalTopic(instruction_unit_id=stored_unit.id, topic_code=topic.code, topic_title=topic.title, description=topic.description, section="Knowledge", asset_type="PPT Videos & Podcast"))
        for topic in unit.skills.case_study_word_document:
            db.add(CaseStudy(instruction_unit_id=stored_unit.id, case_code=topic.code, case_title=topic.title, description=topic.description))
            db.add(CaseStudyDocument(instruction_unit_id=stored_unit.id, case_code=topic.code, case_title=topic.title, description=topic.description))
        for topic in unit.skills.case_study_ppt:
            db.add(CaseStudyPPTTopic(instruction_unit_id=stored_unit.id, topic_code=topic.code, topic_title=topic.title, description=topic.description))
        for topic in unit.skills.case_study_demo_videos:
            db.add(CaseStudyVideoTopic(instruction_unit_id=stored_unit.id, topic_code=topic.code, topic_title=topic.title, description=topic.description))
        for topic in unit.skills.case_study_assignment:
            db.add(SkillsActivity(instruction_unit_id=stored_unit.id, activity_code=topic.code, activity_title=topic.title, description=topic.description, activity_type="Case Study Assignment"))
            db.add(CaseStudyAssignment(instruction_unit_id=stored_unit.id, assignment_code=topic.code, assignment_title=topic.title, description=topic.description))
        for topic in unit.assessment_blueprint.mcq_assessment:
            db.add(MCQAssessment(instruction_unit_id=stored_unit.id, assessment_code=topic.code, assessment_title=topic.title, description=topic.description))
        for topic in unit.assessment_blueprint.assessment_assignment:
            db.add(AssessmentAssignment(instruction_unit_id=stored_unit.id, assignment_code=topic.code, assignment_title=topic.title, description=topic.description))
        for topic in unit.assessment_blueprint.marking_rubrics:
            db.add(MarkingRubric(instruction_unit_id=stored_unit.id, rubric_code=topic.code, rubric_title=topic.title, description=topic.description))
    db.commit()


def _unit_download_text(product: LearningProduct, unit: InstructionUnit, db: Session) -> str:
    blueprint = db.query(ContentBlueprint).filter(ContentBlueprint.instruction_unit_id == unit.id).first()
    assessment = db.query(Assessment).filter(Assessment.instruction_unit_id == unit.id).first()
    knowledge_docs = db.query(KnowledgeDocument).filter(KnowledgeDocument.instruction_unit_id == unit.id).all()
    ppt_topics = db.query(KnowledgePPTTopic).filter(KnowledgePPTTopic.instruction_unit_id == unit.id).all()
    video_topics = db.query(KnowledgeVideoTopic).filter(KnowledgeVideoTopic.instruction_unit_id == unit.id).all()
    activities = db.query(LearningActivity).filter(LearningActivity.instruction_unit_id == unit.id).all()
    cases = db.query(CaseStudyDocument).filter(CaseStudyDocument.instruction_unit_id == unit.id).all()
    case_ppts = db.query(CaseStudyPPTTopic).filter(CaseStudyPPTTopic.instruction_unit_id == unit.id).all()
    case_videos = db.query(CaseStudyVideoTopic).filter(CaseStudyVideoTopic.instruction_unit_id == unit.id).all()
    case_assignments = db.query(CaseStudyAssignment).filter(CaseStudyAssignment.instruction_unit_id == unit.id).all()
    mcqs = db.query(MCQAssessment).filter(MCQAssessment.instruction_unit_id == unit.id).all()
    assessment_assignments = db.query(AssessmentAssignment).filter(AssessmentAssignment.instruction_unit_id == unit.id).all()
    rubrics = db.query(MarkingRubric).filter(MarkingRubric.instruction_unit_id == unit.id).all()

    lines = [
        f"{unit.iu_code} - {unit.title}",
        f"Module Title: {product.title}",
        f"Course: {product.course_name}",
        f"Learning Objective: {unit.learning_goal}",
        f"Estimated Learner Hours: {unit.estimated_hours}",
        "",
        "01. KNOWLEDGE",
        *_topic_lines("Instructional Content Text", knowledge_docs, "topic_code", "topic_title"),
        *_topic_lines("PPT Text", ppt_topics, "topic_code", "topic_title"),
        *_topic_lines("PPT Videos / Podcast", video_topics, "topic_code", "topic_title"),
        *_topic_lines("E-learning Activities", activities, "activity_code", "activity_title"),
        "",
        "02. SKILLS",
        *_topic_lines("Case Study Word Document", cases, "case_code", "case_title"),
        *_topic_lines("Case Study PPT", case_ppts, "topic_code", "topic_title"),
        *_topic_lines("Case Study Demo Videos", case_videos, "topic_code", "topic_title"),
        *_topic_lines("Case Study Assignment", case_assignments, "assignment_code", "assignment_title"),
        "",
        "03. ASSESSMENT",
        *_topic_lines("MCQ Assessment", mcqs, "assessment_code", "assessment_title"),
        *_list_lines("Quiz", _json_load_list(getattr(assessment, "quizzes", "[]") if assessment else "[]")),
        *_topic_lines("Assessment Assignment", assessment_assignments, "assignment_code", "assignment_title"),
        *_topic_lines("Marking Rubrics", rubrics, "rubric_code", "rubric_title"),
        "",
        "Production Notes",
        f"Production Estimate: {blueprint.production_estimate if blueprint else ''}",
        f"Readiness Insight: {blueprint.readiness_insight if blueprint else ''}",
    ]
    return "\n".join(str(line) for line in lines if line is not None)


def _assignment_rows_for_unit(product: LearningProduct, unit: InstructionUnit, db: Session) -> list[dict]:
    saved_assignments = {
        (row.asset_type, row.topic_code): row
        for row in db.query(ContentAssignment).filter(ContentAssignment.product_id == product.id, ContentAssignment.iu_code == unit.iu_code).all()
    }
    source_rows = [
        *_topic_assignment_rows(product, unit, db.query(KnowledgeDocument).filter(KnowledgeDocument.instruction_unit_id == unit.id).all(), "Instructional Text", "topic_code", "topic_title"),
        *_topic_assignment_rows(product, unit, db.query(KnowledgePPTTopic).filter(KnowledgePPTTopic.instruction_unit_id == unit.id).all(), "PPT Slides", "topic_code", "topic_title"),
        *_topic_assignment_rows(product, unit, db.query(KnowledgeVideoTopic).filter(KnowledgeVideoTopic.instruction_unit_id == unit.id).all(), "PPT Video", "topic_code", "topic_title"),
        *_topic_assignment_rows(product, unit, db.query(LearningActivity).filter(LearningActivity.instruction_unit_id == unit.id).all(), "E-learning Activity", "activity_code", "activity_title"),
        *_topic_assignment_rows(product, unit, db.query(CaseStudyDocument).filter(CaseStudyDocument.instruction_unit_id == unit.id).all(), "Case Study Word", "case_code", "case_title"),
        *_topic_assignment_rows(product, unit, db.query(CaseStudyPPTTopic).filter(CaseStudyPPTTopic.instruction_unit_id == unit.id).all(), "Case Study PPT", "topic_code", "topic_title"),
        *_topic_assignment_rows(product, unit, db.query(CaseStudyVideoTopic).filter(CaseStudyVideoTopic.instruction_unit_id == unit.id).all(), "Demo Video", "topic_code", "topic_title"),
        *_topic_assignment_rows(product, unit, db.query(CaseStudyAssignment).filter(CaseStudyAssignment.instruction_unit_id == unit.id).all(), "Case Study Assignment", "assignment_code", "assignment_title"),
        *_topic_assignment_rows(product, unit, db.query(MCQAssessment).filter(MCQAssessment.instruction_unit_id == unit.id).all(), "MCQ Assessment", "assessment_code", "assessment_title"),
        *_topic_assignment_rows(product, unit, db.query(AssessmentAssignment).filter(AssessmentAssignment.instruction_unit_id == unit.id).all(), "Assessment Assignment", "assignment_code", "assignment_title"),
        *_topic_assignment_rows(product, unit, db.query(MarkingRubric).filter(MarkingRubric.instruction_unit_id == unit.id).all(), "Marking Rubric", "rubric_code", "rubric_title"),
    ]
    for row in source_rows:
        saved = saved_assignments.get((row["asset_type"], row["topic_code"]))
        if saved:
            row.update(
                {
                    "owner": saved.owner,
                    "owner_email": saved.owner_email,
                    "start_date": saved.start_date,
                    "deadline": saved.deadline,
                    "status": saved.status,
                    "email_status": saved.email_status,
                    "email_sent_at": saved.email_sent_at,
                    "reminder_sent_at": saved.reminder_sent_at,
                }
            )
    return source_rows


def _topic_assignment_rows(product: LearningProduct, unit: InstructionUnit, rows: list, asset_type: str, code_attr: str, title_attr: str) -> list[dict]:
    return [
        {
            "product_id": product.id,
            "product_title": product.title,
            "iu_code": unit.iu_code,
            "iu_title": unit.title,
            "asset_type": asset_type,
            "topic_code": getattr(row, code_attr, ""),
            "topic_title": getattr(row, title_attr, ""),
            "owner": "",
            "owner_email": "",
            "start_date": "",
            "deadline": "",
            "status": "Pending",
            "email_status": "Not Sent",
            "email_sent_at": "",
            "reminder_sent_at": "",
        }
        for row in rows
    ]


def _ensure_assignment_email_columns(db: Session) -> None:
    existing = {row[1] for row in db.execute(text("PRAGMA table_info(content_assignments)")).fetchall()}
    columns = {
        "owner_email": "VARCHAR(255) DEFAULT ''",
        "email_status": "VARCHAR(80) DEFAULT 'Not Sent'",
        "email_sent_at": "VARCHAR(80) DEFAULT ''",
        "reminder_sent_at": "VARCHAR(80) DEFAULT ''",
    }
    for name, definition in columns.items():
        if name not in existing:
            db.execute(text(f"ALTER TABLE content_assignments ADD COLUMN {name} {definition}"))
    db.commit()


def _ensure_learning_product_columns(db: Session) -> None:
    existing = {row[1] for row in db.execute(text("PRAGMA table_info(learning_products)")).fetchall()}
    columns = {
        "module_code": "VARCHAR(40) DEFAULT ''",
    }
    for name, definition in columns.items():
        if name not in existing:
            db.execute(text(f"ALTER TABLE learning_products ADD COLUMN {name} {definition}"))
    db.commit()


def _upsert_assignment_rows(db: Session, rows: list[dict]) -> int:
    saved = 0
    for row in rows:
        if not isinstance(row, dict):
            continue
        product_id = int(row.get("product_id") or 0)
        iu_code = str(row.get("iu_code") or "")
        topic_code = str(row.get("topic_code") or "")
        asset_type = str(row.get("asset_type") or "")
        if not product_id or not iu_code or not topic_code:
            continue
        unit = db.query(InstructionUnit).filter(InstructionUnit.product_id == product_id, InstructionUnit.iu_code == iu_code).first()
        assignment = (
            db.query(ContentAssignment)
            .filter(
                ContentAssignment.product_id == product_id,
                ContentAssignment.iu_code == iu_code,
                ContentAssignment.topic_code == topic_code,
                ContentAssignment.asset_type == asset_type,
            )
            .first()
        )
        if not assignment:
            assignment = ContentAssignment(product_id=product_id, instruction_unit_id=unit.id if unit else None)
            db.add(assignment)
        assignment.iu_code = iu_code
        assignment.iu_title = str(row.get("iu_title") or (unit.title if unit else ""))
        assignment.asset_type = asset_type
        assignment.topic_code = topic_code
        assignment.topic_title = str(row.get("topic_title") or "")
        assignment.owner = str(row.get("owner") or "")
        assignment.owner_email = str(row.get("owner_email") or "")
        assignment.start_date = str(row.get("start_date") or "")
        assignment.deadline = str(row.get("deadline") or "")
        assignment.status = str(row.get("status") or "Pending")
        assignment.email_status = str(row.get("email_status") or assignment.email_status or "Not Sent")
        assignment.email_sent_at = str(row.get("email_sent_at") or assignment.email_sent_at or "")
        assignment.reminder_sent_at = str(row.get("reminder_sent_at") or assignment.reminder_sent_at or "")
        saved += 1
    return saved


def _assignment_product_unit(payload: dict, db: Session) -> tuple[LearningProduct, InstructionUnit]:
    product_id = int(payload.get("product_id") or 0)
    iu_code = str(payload.get("iu_code") or "").upper()
    product = db.query(LearningProduct).filter(LearningProduct.id == product_id).first()
    if not product:
        raise HTTPException(status_code=404, detail="Product Plan record not found.")
    unit = db.query(InstructionUnit).filter(InstructionUnit.product_id == product.id, InstructionUnit.iu_code == iu_code).first()
    if not unit:
        raise HTTPException(status_code=404, detail=f"{iu_code or 'IU'} was not found for this Product Plan.")
    return product, unit


def _assignment_query(product_id: int, iu_code: str, db: Session) -> list[ContentAssignment]:
    return (
        db.query(ContentAssignment)
        .filter(ContentAssignment.product_id == product_id, ContentAssignment.iu_code == iu_code)
        .order_by(ContentAssignment.asset_type.asc(), ContentAssignment.topic_code.asc())
        .all()
    )


def _assignment_recipients(raw_recipients, assignments: list[ContentAssignment]) -> list[str]:
    recipients: list[str] = []
    source_values = raw_recipients if isinstance(raw_recipients, list) else [raw_recipients or ""]
    source_values.extend(row.owner_email for row in assignments if row.owner_email)
    for value in source_values:
        for address in str(value or "").replace(";", ",").split(","):
            address = address.strip()
            if address and address not in recipients:
                recipients.append(address)
    return recipients


def _content_chat_response(product: LearningProduct, question: str, db: Session) -> dict:
    units = db.query(InstructionUnit).filter(InstructionUnit.product_id == product.id).order_by(InstructionUnit.id.asc()).all()
    q = question.lower()
    fetch_code = _extract_fetch_iu_code(question)
    if fetch_code and any(word in q for word in ["fetch", "generate", "create", "build", "load"]):
        generated = _generate_unit_blueprint_for_product(product, fetch_code, db)
        return {
            "answer": (
                f"{generated.iu_code} has been fetched and saved. "
                f"LIA generated the IU blueprint for {generated.title}, including knowledge topics, skills assets, assessment assignment, quizzes, MCQs, and marking rubrics. "
                "You can now open the IU Breakdown, Content Blueprint, Assessment Planning, or Project Brief tabs and use this IU as saved context."
            ),
            "suggested_actions": [
                f"Review {generated.iu_code} in the IU Breakdown Viewer.",
                "Click Approve & Save Review after checking the generated IU content.",
                "Ask LIA Chat about project tasks, knowledge coverage, skills evidence, or rubrics for the fetched IU.",
            ],
            "confidence": "High",
            "fetched_iu": generated,
        }
    context = _content_chat_context(product, units, db)
    if any(word in q for word in ["rubric", "mark", "grade", "assessment"]):
        answer = (
            f"For {product.title}, assessment evidence should be checked against the saved IU assessment assignments and marking rubrics. "
            f"The strongest rubric coverage is: {_sentence_list(context['rubrics'][:5])}. "
            "Use these criteria to keep grading tied to observable learner evidence rather than generic completion."
        )
        actions = [
            "Open Assessment Planning and confirm each IU has an assessment assignment plus marking rubrics.",
            "Download the Project Brief after saving so the rubric section reflects the latest reviewed IU data.",
            "Check that every task asks learners to submit evidence that can be matched to a rubric criterion.",
        ]
    elif any(word in q for word in ["project", "brief", "task", "scenario", "capstone"]):
        answer = (
            f"The project brief should use {product.title} as an integrated workplace scenario. "
            f"It can be structured around {len(units)} IU task areas: {_sentence_list([f'{unit.iu_code} - {unit.title}' for unit in units])}. "
            "Each task should combine the saved knowledge topics with the skills or assessment assignment from the same IU."
        )
        actions = [
            "Review each IU blueprint before approving the product plan.",
            "Click Approve & Save Review so the downloadable brief uses the final IU knowledge and skills.",
            "Use the Project Brief Generator download buttons to produce the DOCX or PPT version.",
        ]
    elif any(word in q for word in ["knowledge", "topic", "content", "ppt", "video"]):
        answer = (
            f"The key knowledge coverage currently comes from these saved production topics: {_sentence_list(context['knowledge'][:7])}. "
            "These should feed the project objectives and learner guidance before learners attempt applied tasks."
        )
        actions = [
            "Fetch any IU that still shows Ready so its knowledge topics are generated.",
            "Review instructional text, PPT, video, and e-learning topics for overlap.",
            "Save the review after editing the IU knowledge sections.",
        ]
    elif any(word in q for word in ["skill", "case", "practice", "assignment", "activity"]):
        answer = (
            f"The strongest skills evidence should come from these IU assignments and case-study assets: {_sentence_list(context['skills'][:7])}. "
            "Use them as the basis for learner deliverables, project tasks, and evidence requirements."
        )
        actions = [
            "Check the Skills sections for case study word, PPT, demo video, and case study assignment coverage.",
            "Make sure each project task has a clear learner deliverable.",
            "Use the assignment board to allocate owners and deadlines for each production asset.",
        ]
    elif any(word in q for word in ["owner", "deadline", "handoff", "status", "production"]):
        answer = (
            f"The production handoff has {context['assignment_count']} assignable topic rows across {len(units)} IUs. "
            f"Current active task signals: {_sentence_list(context['assignment_status'][:6])}. "
            "Prioritize rows without owners or deadlines before sending IU handoff emails."
        )
        actions = [
            "Open Production Handoff and fill owner, owner email, start date, deadline, and status.",
            "Save Assignments before sending IU handoff emails.",
            "Use reminders only for tasks that are not completed.",
        ]
    else:
        answer = (
            f"LIA Content Agentic AI is reading the saved Product Plan for {product.title}. "
            f"There are {len(units)} reviewed IUs, with knowledge, skills, assessment, project brief, and production handoff data available where generated. "
            "Ask about project brief tasks, IU knowledge, skills evidence, marking rubrics, or production handoff priorities."
        )
        actions = [
            "Fetch missing IU blueprints before asking for detailed task or rubric guidance.",
            "Approve and save the review after edits so chat and downloads use the same source data.",
            "Use specific prompts such as 'What tasks should be in the project brief?' or 'Which rubrics need attention?'",
        ]
    confidence = "High" if units and (context["knowledge"] or context["skills"] or context["rubrics"]) else "Medium"
    return {"answer": answer, "suggested_actions": actions, "confidence": confidence}


def _extract_fetch_iu_code(question: str) -> str:
    import re

    match = re.search(r"\bIU\s*0?([1-9][0-9]?)\b", question or "", flags=re.IGNORECASE)
    if not match:
        return ""
    return f"IU{int(match.group(1)):02d}"


def _content_chat_context(product: LearningProduct, units: list[InstructionUnit], db: Session) -> dict:
    knowledge: list[str] = []
    skills: list[str] = []
    rubrics: list[str] = []
    assignment_status: list[str] = []
    assignment_count = 0
    for unit in units:
        knowledge.extend(f"{unit.iu_code}: {item}" for item in _brief_topic_rows(db.query(KnowledgeDocument).filter(KnowledgeDocument.instruction_unit_id == unit.id).all(), "topic_code", "topic_title", "Instructional Text"))
        knowledge.extend(f"{unit.iu_code}: {item}" for item in _brief_topic_rows(db.query(KnowledgePPTTopic).filter(KnowledgePPTTopic.instruction_unit_id == unit.id).all(), "topic_code", "topic_title", "PPT Slides"))
        knowledge.extend(f"{unit.iu_code}: {item}" for item in _brief_topic_rows(db.query(KnowledgeVideoTopic).filter(KnowledgeVideoTopic.instruction_unit_id == unit.id).all(), "topic_code", "topic_title", "PPT Video"))
        skills.extend(f"{unit.iu_code}: {item}" for item in _brief_topic_rows(db.query(CaseStudyDocument).filter(CaseStudyDocument.instruction_unit_id == unit.id).all(), "case_code", "case_title", "Case Study Word"))
        skills.extend(f"{unit.iu_code}: {item}" for item in _brief_topic_rows(db.query(CaseStudyAssignment).filter(CaseStudyAssignment.instruction_unit_id == unit.id).all(), "assignment_code", "assignment_title", "Case Study Assignment"))
        skills.extend(f"{unit.iu_code}: {item}" for item in _brief_topic_rows(db.query(AssessmentAssignment).filter(AssessmentAssignment.instruction_unit_id == unit.id).all(), "assignment_code", "assignment_title", "Assessment Assignment"))
        rubrics.extend(f"{unit.iu_code}: {item}" for item in _brief_topic_rows(db.query(MarkingRubric).filter(MarkingRubric.instruction_unit_id == unit.id).all(), "rubric_code", "rubric_title", "Marking Rubric"))
        rows = _assignment_rows_for_unit(product, unit, db)
        assignment_count += len(rows)
        missing_owner = len([row for row in rows if not (row.get("owner") or "").strip()])
        missing_deadline = len([row for row in rows if not (row.get("deadline") or "").strip()])
        if rows:
            assignment_status.append(f"{unit.iu_code}: {missing_owner} missing owners, {missing_deadline} missing deadlines")
    return {
        "knowledge": knowledge,
        "skills": skills,
        "rubrics": rubrics,
        "assignment_status": assignment_status,
        "assignment_count": assignment_count,
    }


def _sentence_list(items: list[str]) -> str:
    clean = [str(item).strip() for item in items if str(item).strip()]
    if not clean:
        return "no saved items yet"
    return "; ".join(clean[:8])


def _assignment_email_body(product: LearningProduct, unit: InstructionUnit, assignments: list[ContentAssignment], reminder: bool) -> str:
    rows = assignments[:60]
    task_lines = [
        f"- {row.asset_type}: {row.topic_code} - {row.topic_title} | Owner: {row.owner or 'TBC'} | Start: {row.start_date or 'TBC'} | Deadline: {row.deadline or 'TBC'} | Status: {row.status or 'Pending'}"
        for row in rows
    ]
    if reminder:
        opening = (
            "This is a gentle follow-up on the content production tasks that are still not marked as completed. "
            "Please review your assigned items, update the status, and flag any blockers early so we can keep the module delivery on track."
        )
    else:
        opening = (
            "Please find the IU content production handoff for review and implementation. "
            "The attached file contains the IU blueprint, while the task list below shows the current assignment plan."
        )
    return "\n".join(
        [
            "Hi Team,",
            "",
            opening,
            "",
            f"Module: {product.title}",
            f"Instruction Unit: {unit.iu_code} - {unit.title}",
            f"Learning Objective: {unit.learning_goal or 'Review the attached IU blueprint.'}",
            "",
            "Assigned production tasks:",
            *(task_lines or ["- No assigned production tasks were found."]),
            "",
            "Please update the assignment board once your content is completed or ready for review.",
            "",
            "Warm regards,",
            "EI THANDAR KHAING, Learning Management",
            "e: eithandar@educlaas.com",
            "o: +959 798469301",
            "w: www.claas2saas.com",
            "",
            "a: 11 Eunos Road 8, #07-02 Lifelong Learning Institute, Singapore 408601",
        ]
    )


def _unit_download_attachment(product: LearningProduct, unit: InstructionUnit, selected_format: str, db: Session) -> dict:
    content = _unit_download_text(product, unit, db)
    safe_name = f"{unit.iu_code}-{_download_safe_name(unit.title)}"
    if selected_format == "pdf":
        return {"filename": f"{safe_name}.pdf", "content": _simple_pdf(content), "maintype": "application", "subtype": "pdf"}
    if selected_format in {"xls", "xlsx"}:
        return {"filename": f"{safe_name}.xls", "content": _unit_xls(content, unit), "maintype": "application", "subtype": "vnd.ms-excel"}
    return {
        "filename": f"{safe_name}.docx",
        "content": _unit_docx(content),
        "maintype": "application",
        "subtype": "vnd.openxmlformats-officedocument.wordprocessingml.document",
    }


def _project_brief_download_text(product: LearningProduct, db: Session) -> dict:
    units = db.query(InstructionUnit).filter(InstructionUnit.product_id == product.id).order_by(InstructionUnit.id.asc()).all()
    project = db.query(ProjectBrief).filter(ProjectBrief.product_id == product.id).first()
    learning_outcomes = _split_learning_outcomes(product.learning_outcomes)
    if not learning_outcomes:
        learning_outcomes = [unit.learning_goal for unit in units if unit.learning_goal]

    tasks: list[dict] = []
    rubric_items: list[str] = []
    knowledge_items: list[str] = []
    skills_items: list[str] = []
    for unit in units:
        knowledge_rows = [
            *_brief_topic_rows(db.query(KnowledgeDocument).filter(KnowledgeDocument.instruction_unit_id == unit.id).all(), "topic_code", "topic_title", "Instructional Text"),
            *_brief_topic_rows(db.query(KnowledgePPTTopic).filter(KnowledgePPTTopic.instruction_unit_id == unit.id).all(), "topic_code", "topic_title", "PPT Slides"),
            *_brief_topic_rows(db.query(KnowledgeVideoTopic).filter(KnowledgeVideoTopic.instruction_unit_id == unit.id).all(), "topic_code", "topic_title", "PPT Video"),
            *_brief_topic_rows(db.query(LearningActivity).filter(LearningActivity.instruction_unit_id == unit.id).all(), "activity_code", "activity_title", "E-learning Activity"),
        ]
        skills_rows = [
            *_brief_topic_rows(db.query(CaseStudyDocument).filter(CaseStudyDocument.instruction_unit_id == unit.id).all(), "case_code", "case_title", "Case Study Word"),
            *_brief_topic_rows(db.query(CaseStudyPPTTopic).filter(CaseStudyPPTTopic.instruction_unit_id == unit.id).all(), "topic_code", "topic_title", "Case Study PPT"),
            *_brief_topic_rows(db.query(CaseStudyVideoTopic).filter(CaseStudyVideoTopic.instruction_unit_id == unit.id).all(), "topic_code", "topic_title", "Demo Video"),
            *_brief_topic_rows(db.query(CaseStudyAssignment).filter(CaseStudyAssignment.instruction_unit_id == unit.id).all(), "assignment_code", "assignment_title", "Case Study Assignment"),
        ]
        assessment_rows = _brief_topic_rows(
            db.query(AssessmentAssignment).filter(AssessmentAssignment.instruction_unit_id == unit.id).all(),
            "assignment_code",
            "assignment_title",
            "Assessment Assignment",
        )
        rubrics = _brief_topic_rows(
            db.query(MarkingRubric).filter(MarkingRubric.instruction_unit_id == unit.id).all(),
            "rubric_code",
            "rubric_title",
            "Marking Rubric",
        )

        knowledge_items.extend(f"{unit.iu_code} - {item}" for item in knowledge_rows)
        skills_items.extend(f"{unit.iu_code} - {item}" for item in skills_rows)
        rubric_items.extend(f"{unit.iu_code} - {item}" for item in rubrics)
        task_sources = assessment_rows or [item for item in skills_rows if "Assignment" in item] or skills_rows[:2] or knowledge_rows[:2]
        for item in task_sources:
            tasks.append(
                {
                    "title": f"{unit.iu_code}: {unit.title}",
                    "objective": unit.learning_goal or f"Apply the knowledge and skills from {unit.iu_code}.",
                    "description": item,
                    "knowledge": knowledge_rows[:4],
                    "skills": skills_rows[:4],
                }
            )

    if not tasks:
        tasks = [
            {
                "title": "Task 1: Integrated Project Submission",
                "objective": "Apply the reviewed IU knowledge and skills in one practical project.",
                "description": "Prepare a project submission that demonstrates understanding, applied practice, and evidence of learning.",
                "knowledge": knowledge_items[:4],
                "skills": skills_items[:4],
            }
        ]

    return {
        "title": f"{product.title} Project Brief",
        "module": product.title,
        "course": product.course_name,
        "project_scenario": (project.capstone_scenario if project else "") or f"Learners will complete an applied workplace scenario for {product.title}, using the saved IU knowledge and skills assignments as evidence of competence.",
        "project_objectives": (project.project_deliverables if project else "[]"),
        "learning_outcomes": learning_outcomes,
        "project_format": ["DOCX written brief with evidence links", "PPT presentation summary for project walkthrough"],
        "tasks": tasks,
        "marking_rubrics": rubric_items or _json_load_list(project.evaluation_criteria if project else "[]"),
        "knowledge_items": knowledge_items,
        "skills_items": skills_items,
    }


def _brief_topic_rows(rows: list, code_attr: str, title_attr: str, label: str) -> list[str]:
    output = []
    for row in rows:
        code = getattr(row, code_attr, "")
        title = getattr(row, title_attr, "")
        description = getattr(row, "description", "")
        prefix = f"{label}: "
        topic = " - ".join(part for part in [code, title] if part)
        if description:
            topic = f"{topic}: {description}" if topic else description
        if topic:
            output.append(f"{prefix}{topic}")
    return output


def _split_learning_outcomes(value: str) -> list[str]:
    if not value:
        return []
    separators = ["\n", ";"]
    items = [value]
    for separator in separators:
        items = [piece for item in items for piece in item.split(separator)]
    cleaned = []
    for item in items:
        text = item.strip(" -\t\r")
        if text and text not in cleaned:
            cleaned.append(text)
    return cleaned


def _project_brief_docx(content: dict) -> bytes:
    document = Document()
    document.add_heading(content["title"], level=1)
    document.add_paragraph(f"Module: {content['module']}")
    document.add_paragraph(f"Course: {content['course']}")

    document.add_heading("Project Scenario", level=2)
    document.add_paragraph(content["project_scenario"])

    document.add_heading("Project Objectives", level=2)
    for item in _json_load_list(content["project_objectives"]) or content["skills_items"][:5]:
        document.add_paragraph(str(item), style="List Bullet")

    document.add_heading("Learning Outcomes", level=2)
    for item in content["learning_outcomes"] or ["Learning outcomes will be confirmed from the saved Product Plan review."]:
        document.add_paragraph(str(item), style="List Bullet")

    document.add_heading("Project Format - DOCX / PPT", level=2)
    for item in content["project_format"]:
        document.add_paragraph(item, style="List Bullet")

    document.add_heading("Tasks", level=2)
    for index, task in enumerate(content["tasks"], start=1):
        document.add_heading(f"Task {index}: {task['title']}", level=3)
        document.add_paragraph(task["description"])
        document.add_paragraph(f"Objective: {task['objective']}")
        for item in task["knowledge"]:
            document.add_paragraph(f"Knowledge: {item}", style="List Bullet")
        for item in task["skills"]:
            document.add_paragraph(f"Skill: {item}", style="List Bullet")

    document.add_heading("Marking Rubrics - Details", level=2)
    for item in content["marking_rubrics"] or ["Rubrics will be finalized after IU assessment planning is reviewed and saved."]:
        document.add_paragraph(str(item), style="List Bullet")

    buffer = BytesIO()
    document.save(buffer)
    return buffer.getvalue()


def _project_brief_ppt(content: dict) -> str:
    slides = [
        ("Project Scenario", [content["project_scenario"]]),
        ("Project Objectives", _json_load_list(content["project_objectives"]) or content["skills_items"][:5]),
        ("Learning Outcomes", content["learning_outcomes"]),
        ("Project Format - DOCX / PPT", content["project_format"]),
        ("Tasks", [f"Task {index}: {task['title']} - {task['description']}" for index, task in enumerate(content["tasks"], start=1)]),
        ("Marking Rubrics - Details", content["marking_rubrics"]),
    ]
    sections = []
    for title, items in slides:
        bullets = "".join(f"<li>{_html_escape(item)}</li>" for item in (items or ["Pending review."])[:12])
        sections.append(
            "<section style='page-break-after:always;width:960px;height:540px;padding:42px;font-family:Arial,sans-serif;'>"
            f"<h1>{_html_escape(title)}</h1><ul style='font-size:24px;line-height:1.35;'>{bullets}</ul></section>"
        )
    return f"<html><body><h1>{_html_escape(content['title'])}</h1>{''.join(sections)}</body></html>"


def _now_stamp() -> str:
    return datetime.now(timezone.utc).isoformat(timespec="seconds")


def _topic_lines(title: str, rows: list, code_attr: str, title_attr: str) -> list[str]:
    lines = [f"\n{title}"]
    if not rows:
        lines.append("Content pending.")
        return lines
    for row in rows:
        code = getattr(row, code_attr, "")
        topic_title = getattr(row, title_attr, "")
        description = getattr(row, "description", "")
        lines.extend([f"{code} - {topic_title}", f"Description: {description}", ""])
    return lines


def _list_lines(title: str, items: list[str]) -> list[str]:
    lines = [f"\n{title}"]
    if not items:
        lines.append("Content pending.")
        return lines
    for item in items:
        lines.extend([str(item), ""])
    return lines


def _json_load_list(value: str) -> list[str]:
    try:
        parsed = json.loads(value or "[]")
        return parsed if isinstance(parsed, list) else []
    except json.JSONDecodeError:
        return []


def _normalize_lookup_text(value: str) -> str:
    return " ".join((value or "").lower().split())


def _download_safe_name(value: str) -> str:
    import re

    clean = re.sub(r"[^A-Za-z0-9_-]+", "-", value.strip()).strip("-")
    return clean[:70] or "instruction-unit"


def _unit_docx(content: str) -> bytes:
    document = Document()
    for index, line in enumerate(content.splitlines()):
        if index == 0:
            document.add_heading(line, level=1)
        elif line.startswith(("01.", "02.", "03.")):
            document.add_heading(line, level=2)
        elif line.strip():
            document.add_paragraph(line)
    buffer = BytesIO()
    document.save(buffer)
    return buffer.getvalue()


def _unit_xls(content: str, unit: InstructionUnit) -> str:
    rows = []
    section = ""
    for line in content.splitlines():
        clean = line.strip()
        if not clean:
            continue
        if clean.startswith(("01.", "02.", "03.")):
            section = clean
            continue
        rows.append((section, clean))
    cells = "\n".join(
        f"<tr><td>{_html_escape(unit.iu_code)}</td><td>{_html_escape(section)}</td><td>{_html_escape(item)}</td></tr>"
        for section, item in rows
    )
    return f"<html><body><table border='1'><tr><th>IU</th><th>Section</th><th>Content</th></tr>{cells}</table></body></html>"


def _simple_pdf(content: str) -> bytes:
    lines = [line[:95] for line in content.splitlines()[:70]]
    stream_lines = ["BT", "/F1 10 Tf", "50 780 Td"]
    first = True
    for line in lines:
        escaped = line.replace("\\", "\\\\").replace("(", "\\(").replace(")", "\\)")
        if first:
            stream_lines.append(f"({escaped}) Tj")
            first = False
        else:
            stream_lines.append(f"0 -14 Td ({escaped}) Tj")
    stream_lines.append("ET")
    stream = "\n".join(stream_lines).encode("latin-1", errors="ignore")
    objects = [
        b"1 0 obj << /Type /Catalog /Pages 2 0 R >> endobj",
        b"2 0 obj << /Type /Pages /Kids [3 0 R] /Count 1 >> endobj",
        b"3 0 obj << /Type /Page /Parent 2 0 R /MediaBox [0 0 612 792] /Resources << /Font << /F1 4 0 R >> >> /Contents 5 0 R >> endobj",
        b"4 0 obj << /Type /Font /Subtype /Type1 /BaseFont /Helvetica >> endobj",
        b"5 0 obj << /Length " + str(len(stream)).encode() + b" >> stream\n" + stream + b"\nendstream endobj",
    ]
    pdf = bytearray(b"%PDF-1.4\n")
    offsets = [0]
    for obj in objects:
        offsets.append(len(pdf))
        pdf.extend(obj + b"\n")
    xref_offset = len(pdf)
    pdf.extend(f"xref\n0 {len(objects) + 1}\n0000000000 65535 f \n".encode())
    for offset in offsets[1:]:
        pdf.extend(f"{offset:010d} 00000 n \n".encode())
    pdf.extend(f"trailer << /Size {len(objects) + 1} /Root 1 0 R >>\nstartxref\n{xref_offset}\n%%EOF".encode())
    return bytes(pdf)


def _html_escape(value: str) -> str:
    return str(value).replace("&", "&amp;").replace("<", "&lt;").replace(">", "&gt;")


def _extract_product_plan_text(content: bytes, filename: str) -> str:
    suffix = filename.lower().rsplit(".", 1)[-1] if "." in filename else ""
    if suffix == "pdf":
        try:
            from pypdf import PdfReader

            reader = PdfReader(BytesIO(content))
            pages = [(page.extract_text() or "").strip() for page in reader.pages]
            extracted = "\n".join(page for page in pages if page).strip()
            if extracted:
                return extracted
        except Exception:
            pass
        decoded = content.decode("utf-8", errors="ignore")
        return decoded.strip() or f"Uploaded PDF product plan: {filename}. Use product title, module structure, outcomes, hours, and delivery modes supplied by the user."
    if suffix == "docx":
        try:
            document = Document(BytesIO(content))
            paragraphs = [paragraph.text.strip() for paragraph in document.paragraphs if paragraph.text.strip()]
            table_cells = []
            for table in document.tables:
                for row in table.rows:
                    values = [cell.text.strip() for cell in row.cells if cell.text.strip()]
                    if values:
                        table_cells.append(" | ".join(values))
            extracted = "\n".join([*paragraphs, *table_cells]).strip()
            if extracted:
                return extracted
        except Exception:
            pass
    return content.decode("utf-8", errors="ignore").strip()


def _metadata_from_product_plan(text: str, filename: str) -> dict:
    source = text or filename
    title = _first_matching_line(source, ["product plan", "module", "course", "diploma"]) or _clean_filename(filename)
    course_name = _detect_course_name(source, title)
    learner_level = "Advanced" if any(word in source.lower() for word in ["advanced", "higher diploma", "professional diploma"]) else "Basic"
    module_structure = _extract_iu_structure(source)
    learning_outcomes = _extract_learning_outcomes(source)
    delivery_modes = _extract_delivery_modes(source)
    return {
        "title": title,
        "module_code": _extract_module_code(source),
        "course_name": course_name,
        "learner_level": learner_level,
        "module_structure": module_structure,
        "learning_outcomes": learning_outcomes,
        "total_learning_hours": _extract_total_hours(source),
        "delivery_modes": delivery_modes,
    }


def _detect_course_name(text: str, fallback: str) -> str:
    lower = text.lower()
    options = [
        "Foundation Diploma in Application Development",
        "Higher Diploma in Software Engineering",
        "Professional Diploma in Full Stack Web development",
    ]
    for option in options:
        if option.lower() in lower:
            return option
    if "software engineering" in lower:
        return "Higher Diploma in Software Engineering"
    if "full stack" in lower or "web development" in lower:
        return "Professional Diploma in Full Stack Web development"
    if "application development" in lower:
        return "Foundation Diploma in Application Development"
    return fallback or "Uploaded Product Plan"


def _extract_module_code(text: str) -> str:
    import re

    patterns = [
        r"module\s*code\s*[:\-]?\s*([A-Z0-9][A-Z0-9\-_/]{1,18})",
        r"\bcode\s*[:\-]?\s*([A-Z0-9][A-Z0-9\-_/]{1,18})",
        r"module\s*code\s*[:\-]?\s*([A-Z]{1,6}\s*[-_/]?\s*\d{1,4}[A-Z]?)",
        r"module\s*[:\-]?\s*([A-Z]{1,6}\s*[-_/]?\s*\d{1,4}[A-Z]?)",
        r"\b([A-Z]{2,6}\s*[-_/]?\s*\d{2,4}[A-Z]?)\b",
    ]
    for pattern in patterns:
        match = re.search(pattern, text, flags=re.IGNORECASE)
        if match:
            code = re.sub(r"\s+", "", match.group(1).upper()).strip("-_/")
            code = re.split(r"[^A-Z0-9_\-/]", code)[0]
            if 2 <= len(code) <= 18 and code not in {"MODULE", "COURSE", "PRODUCT", "PLAN"}:
                return code
    return ""


def _extract_iu_structure(text: str) -> str:
    lines = [line.strip() for line in text.splitlines() if line.strip()]
    iu_lines = [line for line in lines if line.lower().startswith("iu") and any(ch.isdigit() for ch in line[:6])]
    if len(iu_lines) >= 5:
        return "; ".join(iu_lines[:5])
    return ""


def _extract_learning_outcomes(text: str) -> str:
    lines = [line.strip(" -\t") for line in text.splitlines() if line.strip()]
    outcome_lines = [line for line in lines if any(key in line.lower() for key in ["learning outcome", "lo", "able to", "explain", "apply", "build", "design", "develop"])]
    if outcome_lines:
        return "\n".join(outcome_lines[:8])
    return ""


def _extract_delivery_modes(text: str) -> str:
    lower = text.lower()
    modes = []
    if "self" in lower or "asynchronous" in lower or "async" in lower:
        modes.append("Self-paced")
    if "e-learning" in lower or "elearning" in lower:
        modes.append("E-learning")
    if "assessment" in lower:
        modes.append("Assessment-led")
    if "mentor" in lower:
        modes.append("Limited mentoring support")
    return ", ".join(modes) or ""


def _extract_total_hours(text: str) -> float:
    import re

    lower = text.lower()
    total_match = re.search(r"(?:total|overall|module)\s+(?:learning\s+)?(?:duration|hours|hrs).*?(\d+(?:\.\d+)?)\s*(?:hours|hrs|hr)", lower)
    if total_match:
        try:
            return float(total_match.group(1))
        except ValueError:
            pass

    session_values = []
    for line in lower.splitlines():
        if any(skip in line for skip in ["content development", "ppt", "video", "word", "editing"]):
            continue
        if any(key in line for key in ["session plan", "session", "duration", "learning hours", "iu"]):
            for match in re.findall(r"(\d+(?:\.\d+)?)\s*(?:hours|hrs|hr)", line):
                value = float(match)
                if 0 < value <= 24:
                    session_values.append(value)
    if len(session_values) >= 2:
        return round(sum(session_values), 1)

    matches = re.findall(r"(\d+(?:\.\d+)?)\s*(?:hours|hrs|hr)", lower)
    if matches:
        try:
            return float(matches[0])
        except ValueError:
            pass
    return 0


def _first_matching_line(text: str, keywords: list[str]) -> str:
    for line in text.splitlines():
        clean = line.strip(" -\t")
        if 8 <= len(clean) <= 140 and any(keyword in clean.lower() for keyword in keywords):
            return clean
    return ""


def _clean_filename(filename: str) -> str:
    clean = filename.rsplit(".", 1)[0].replace("_", " ").replace("-", " ").strip()
    return clean or "Uploaded Product Plan"
